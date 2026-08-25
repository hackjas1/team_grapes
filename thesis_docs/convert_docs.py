import sys
import os
import re
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'''
        <w:tcMar {nsdecls("w")}>
            <w:top w:w="{top}" w:type="dxa"/>
            <w:bottom w:w="{bottom}" w:type="dxa"/>
            <w:left w:w="{left}" w:type="dxa"/>
            <w:right w:w="{right}" w:type="dxa"/>
        </w:tcMar>
    ''')
    tcPr.append(tcMar)

def set_table_borders(table, color="D1D5DB", sz="4", val="single"):
    tblPr = table._tbl.tblPr
    tblBorders = parse_xml(f'''
        <w:tblBorders {nsdecls("w")}>
            <w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
            <w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
            <w:left w:val="none"/>
            <w:right w:val="none"/>
            <w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
            <w:insideV w:val="none"/>
        </w:tblBorders>
    ''')
    tblPr.append(tblBorders)

def parse_markdown_table(lines):
    rows = []
    for line in lines:
        line = line.strip()
        if not line or not line.startswith('|'):
            continue
        if re.match(r'^\|[\s\-:|]+\|$', line):
            continue  # separator row
        cells = [c.strip() for c in line.split('|')[1:-1]]
        rows.append(cells)
    return rows

def format_inline_runs(paragraph, text, base_font_size=10.5, is_bold_base=False, is_code_base=False, color_rgb=None):
    # Regex to split on bold, code, math symbols, links
    pattern = r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`|\$[^$]+\$|\[[^\]]+\]\([^)]+\))'
    parts = re.split(pattern, text)
    
    for part in parts:
        if not part:
            continue
        run = paragraph.add_run()
        run.font.name = 'Calibri'
        run.font.size = Pt(base_font_size)
        if color_rgb:
            run.font.color.rgb = color_rgb
        
        if part.startswith('**') and part.endswith('**'):
            run.text = part[2:-2]
            run.bold = True
        elif part.startswith('*') and part.endswith('*'):
            run.text = part[1:-1]
            run.italic = True
        elif part.startswith('`') and part.endswith('`'):
            run.text = part[1:-1]
            run.font.name = 'Consolas'
            run.font.size = Pt(base_font_size - 0.5)
            run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
        elif part.startswith('$') and part.endswith('$'):
            clean_math = part[1:-1].replace('\\text', '').replace('{', '').replace('}', '').replace('\\le', ' ≤ ').replace('\\ge', ' ≥ ').replace('\\times', ' × ').replace('\\parallel', ' ∥ ').replace('\\sin^2', 'sin²').replace('\\cos', 'cos').replace('\\Delta', 'Δ').replace('\\phi', 'φ').replace('\\lambda', 'λ').replace('\\operatorname', '').replace('\\left', '').replace('\\right', '').replace('\\sqrt', '√')
            run.text = clean_math
            run.font.name = 'Cambria Math'
            run.italic = True
        elif part.startswith('[') and ']' in part and '(' in part:
            m = re.match(r'\[([^\]]+)\]\(([^)]+)\)', part)
            if m:
                run.text = m.group(1)
                run.font.color.rgb = RGBColor(0x02, 0x84, 0xC7)
                run.underline = True
            else:
                run.text = part
        else:
            run.text = part
            if is_bold_base:
                run.bold = True
            if is_code_base:
                run.font.name = 'Consolas'

def convert_md_to_docx(md_path, docx_path):
    print(f"Converting {md_path} -> {docx_path}...")
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    doc = Document()
    
    # Set standard page margins (1 inch / 1440 dxa)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    lines = content.split('\n')
    i = 0
    in_code_block = False
    code_block_lines = []
    in_table = False
    table_lines = []

    NAVY = RGBColor(0x0B, 0x20, 0x46)
    BLUE = RGBColor(0x00, 0x70, 0xBA)
    DARK = RGBColor(0x1F, 0x29, 0x37)
    MUTED = RGBColor(0x4B, 0x55, 0x63)

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Handle Code Block
        if stripped.startswith('```'):
            if in_code_block:
                # End of code block -> write box
                in_code_block = False
                code_text = '\n'.join(code_block_lines)
                code_block_lines = []
                
                tbl = doc.add_table(rows=1, cols=1)
                tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
                cell = tbl.cell(0, 0)
                set_cell_background(cell, "F8FAFC")
                set_cell_margins(cell, top=140, bottom=140, left=200, right=200)
                set_table_borders(tbl, color="CBD5E1", sz="6", val="single")
                
                p = cell.paragraphs[0]
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.line_spacing = 1.15
                
                for cl in code_text.split('\n'):
                    format_inline_runs(p, cl + '\n', base_font_size=9.0, is_code_base=True, color_rgb=RGBColor(0x0F, 0x17, 0x2A))
                
                doc.add_paragraph().paragraph_format.space_after = Pt(4)
            else:
                in_code_block = True
                code_block_lines = []
            i += 1
            continue

        if in_code_block:
            code_block_lines.append(line)
            i += 1
            continue

        # Handle Tables
        if stripped.startswith('|') and stripped.endswith('|'):
            if not in_table:
                in_table = True
                table_lines = [line]
            else:
                table_lines.append(line)
            i += 1
            continue
        elif in_table:
            in_table = False
            rows_data = parse_markdown_table(table_lines)
            table_lines = []
            if rows_data:
                num_rows = len(rows_data)
                num_cols = max(len(r) for r in rows_data)
                tbl = doc.add_table(rows=num_rows, cols=num_cols)
                tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
                set_table_borders(tbl, color="CBD5E1", sz="4", val="single")
                
                for r_idx, row in enumerate(rows_data):
                    is_header = (r_idx == 0)
                    for c_idx in range(num_cols):
                        val = row[c_idx] if c_idx < len(row) else ""
                        cell = tbl.cell(r_idx, c_idx)
                        set_cell_margins(cell, top=100, bottom=100, left=140, right=140)
                        
                        if is_header:
                            set_cell_background(cell, "0B2046")
                        elif r_idx % 2 == 1:
                            set_cell_background(cell, "F8FAFC")
                        else:
                            set_cell_background(cell, "FFFFFF")
                            
                        p = cell.paragraphs[0]
                        p.paragraph_format.space_before = Pt(2)
                        p.paragraph_format.space_after = Pt(2)
                        p.paragraph_format.line_spacing = 1.1
                        
                        font_color = RGBColor(0xFF, 0xFF, 0xFF) if is_header else DARK
                        format_inline_runs(p, val, base_font_size=9.0 if not is_header else 9.5, is_bold_base=is_header, color_rgb=font_color)
                        
                p_spacer = doc.add_paragraph()
                p_spacer.paragraph_format.space_after = Pt(6)

        # Skip empty lines
        if not stripped:
            i += 1
            continue

        # Handle Headings
        if stripped.startswith('# '):
            h = doc.add_paragraph()
            h.paragraph_format.space_before = Pt(16)
            h.paragraph_format.space_after = Pt(6)
            h.paragraph_format.keep_with_next = True
            format_inline_runs(h, stripped[2:], base_font_size=17, is_bold_base=True, color_rgb=NAVY)
        elif stripped.startswith('## '):
            h = doc.add_paragraph()
            h.paragraph_format.space_before = Pt(13)
            h.paragraph_format.space_after = Pt(5)
            h.paragraph_format.keep_with_next = True
            format_inline_runs(h, stripped[3:], base_font_size=14, is_bold_base=True, color_rgb=BLUE)
        elif stripped.startswith('### '):
            h = doc.add_paragraph()
            h.paragraph_format.space_before = Pt(10)
            h.paragraph_format.space_after = Pt(4)
            h.paragraph_format.keep_with_next = True
            format_inline_runs(h, stripped[4:], base_font_size=12, is_bold_base=True, color_rgb=NAVY)
        elif stripped.startswith('#### '):
            h = doc.add_paragraph()
            h.paragraph_format.space_before = Pt(8)
            h.paragraph_format.space_after = Pt(3)
            h.paragraph_format.keep_with_next = True
            format_inline_runs(h, stripped[5:], base_font_size=11, is_bold_base=True, color_rgb=DARK)
        elif stripped.startswith('---'):
            # Horizontal line
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            pPr = p._p.get_or_add_pPr()
            pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="6" w:space="1" w:color="CBD5E1"/></w:pBdr>')
            pPr.append(pBdr)
        elif stripped.startswith('* ') or stripped.startswith('- ') or stripped.startswith('• '):
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.15
            content_text = re.sub(r'^[*\-•]\s+', '', stripped)
            format_inline_runs(p, content_text, base_font_size=10.0, color_rgb=DARK)
        elif re.match(r'^\d+\.\s+', stripped):
            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.15
            content_text = re.sub(r'^\d+\.\s+', '', stripped)
            format_inline_runs(p, content_text, base_font_size=10.0, color_rgb=DARK)
        else:
            # Regular paragraph
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.15
            format_inline_runs(p, stripped, base_font_size=10.5, color_rgb=DARK)

        i += 1

    doc.save(docx_path)
    print(f"Successfully generated: {docx_path}")

if __name__ == '__main__':
    base_dir = r"d:\THESIS CAPSTONE SYSTEM\team_grapes_0.4\thesis_docs"
    convert_md_to_docx(os.path.join(base_dir, "SYSTEM_TECHNICAL_MANUAL.md"), os.path.join(base_dir, "SYSTEM_TECHNICAL_MANUAL.docx"))
    convert_md_to_docx(os.path.join(base_dir, "USER_MANUAL.md"), os.path.join(base_dir, "USER_MANUAL.docx"))
